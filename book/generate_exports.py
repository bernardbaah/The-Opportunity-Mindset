#!/usr/bin/env python3
"""
Generate DOCX and print-ready HTML exports for The Opportunity Mindset.
Run from workspace root: python3 book/generate_exports.py
"""

import os
import re
import glob
import io
import base64
import time
import urllib.request
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Image cache ──────────────────────────────────────────────────────────
# Maps URL → (base64_data_uri, raw_bytes) so each image is fetched once
_IMAGE_CACHE = {}

def fetch_image(url):
    """Download an image URL and return (data_uri, bytes). Returns None on failure."""
    if url in _IMAGE_CACHE:
        return _IMAGE_CACHE[url]
    # Use a smaller resolution for Pexels to keep file sizes sane
    url_dl = re.sub(r'w=\d+', 'w=900', url)
    url_dl = re.sub(r'h=\d+', 'h=540', url_dl)
    try:
        req = urllib.request.Request(url_dl, headers={
            'User-Agent': 'Mozilla/5.0 (compatible; BookExporter/1.0)'
        })
        with urllib.request.urlopen(req, timeout=15) as resp:
            raw = resp.read()
        ct = resp.headers.get('Content-Type', 'image/jpeg').split(';')[0].strip()
        if not ct.startswith('image/'):
            ct = 'image/jpeg'
        b64 = base64.b64encode(raw).decode('ascii')
        data_uri = f'data:{ct};base64,{b64}'
        _IMAGE_CACHE[url] = (data_uri, raw)
        return data_uri, raw
    except Exception as e:
        print(f"    [warn] Could not fetch image: {url_dl} — {e}")
        _IMAGE_CACHE[url] = None
        return None

# ── File order ─────────────────────────────────────────────────────────
BOOK_FILES = [
    "book/01_preface.md",
    "book/02_introduction.md",
    "book/03_about.md",
] + [
    f"book/{str(i+4).zfill(2)}_chapter{str(i+1).zfill(2)}.md"
    for i in range(31)
] + [
    "book/35_conclusion.md",
    "book/36_appendix_a.md",
    "book/37_appendix_b.md",
    "book/38_appendix_c.md",
    "book/39_appendix_d.md",
    "book/40_appendix_e.md",
    "book/41_appendix_f.md",
    "book/42_bonus_resources.md",
]

OUTPUT_DIR = Path("book/downloads")
OUTPUT_DIR.mkdir(exist_ok=True)

# ── Helpers ─────────────────────────────────────────────────────────────

def strip_inline(text):
    """Strip markdown inline formatting for plain text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text

def parse_inline(text):
    """Return list of (text, bold, italic) tuples from inline markdown."""
    runs = []
    # pattern matches **bold**, *italic*, ***both***
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|(.+?)(?=\*|\*\*|\*\*\*|`|$))', re.DOTALL)
    remaining = text
    pos = 0
    while pos < len(text):
        # bold+italic
        m = re.match(r'\*\*\*(.+?)\*\*\*', text[pos:], re.DOTALL)
        if m:
            runs.append((m.group(1), True, True))
            pos += m.end(); continue
        # bold
        m = re.match(r'\*\*(.+?)\*\*', text[pos:], re.DOTALL)
        if m:
            runs.append((m.group(1), True, False))
            pos += m.end(); continue
        # italic
        m = re.match(r'\*(.+?)\*', text[pos:], re.DOTALL)
        if m:
            runs.append((m.group(1), False, True))
            pos += m.end(); continue
        # code
        m = re.match(r'`(.+?)`', text[pos:], re.DOTALL)
        if m:
            runs.append((m.group(1), False, False))
            pos += m.end(); continue
        # find next special char
        nxt = len(text)
        for ch in ['***', '**', '*', '`']:
            idx = text.find(ch, pos)
            if idx != -1:
                nxt = min(nxt, idx)
        if nxt > pos:
            runs.append((text[pos:nxt], False, False))
        pos = nxt if nxt > pos else pos + 1
    return runs

def set_run_font(run, bold=False, italic=False, size=12, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph_with_inline(doc, text, style_name="Normal", base_size=12):
    para = doc.add_paragraph(style=style_name)
    # strip markdown links to just text
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    runs = parse_inline(text)
    for txt, bold, italic in runs:
        if not txt:
            continue
        run = para.add_run(txt)
        set_run_font(run, bold=bold, italic=italic, size=base_size)
    return para

def add_heading(doc, text, level):
    """Add a heading, stripping markdown formatting."""
    clean = strip_inline(text).strip()
    if level == 1:
        p = doc.add_heading(clean, level=1)
        p.runs[0].font.size = Pt(22)
        p.runs[0].font.color.rgb = RGBColor(0x1C, 0x25, 0x33)
    elif level == 2:
        p = doc.add_heading(clean, level=2)
        p.runs[0].font.size = Pt(16)
        p.runs[0].font.color.rgb = RGBColor(0x1C, 0x25, 0x33)
    elif level == 3:
        p = doc.add_heading(clean, level=3)
        p.runs[0].font.size = Pt(13)
    else:
        p = doc.add_heading(clean, level=4)
        p.runs[0].font.size = Pt(12)
    return p

def add_table_from_md(doc, rows):
    """Parse markdown table rows and add to document."""
    # rows: list of '| a | b | c |' strings
    # first row = header, second row = separator, rest = data
    def parse_row(line):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        return cells

    header_row = parse_row(rows[0])
    data_rows = [parse_row(r) for r in rows[2:] if r.strip() and not re.match(r'^\|[-| :]+\|$', r.strip())]

    col_count = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    table.style = 'Table Grid'

    # Header
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(header_row):
        if i < len(hdr_cells):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(strip_inline(cell_text))
            run.bold = True
            run.font.size = Pt(10)
            # shade header
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1C2533')
            tcPr.append(shd)
            run.font.color.rgb = RGBColor(0xF0, 0xE8, 0xD8)

    # Data
    for ri, row_data in enumerate(data_rows):
        cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            if ci < len(cells):
                p = cells[ci].paragraphs[0]
                run = p.add_run(strip_inline(cell_text))
                run.font.size = Pt(10)

    doc.add_paragraph()  # spacing

def process_file_to_docx(doc, filepath, is_first=False):
    """Parse a markdown file and add its content to the Word document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    if not is_first:
        doc.add_page_break()

    table_buffer = []
    in_table = False
    in_code = False
    list_items = []

    def flush_list():
        nonlocal list_items
        for item in list_items:
            clean = re.sub(r'^[-*+]\s+', '', item).strip()
            clean = re.sub(r'^\d+\.\s+', '', clean)
            p = add_paragraph_with_inline(doc, clean, base_size=11)
            p.style = doc.styles['List Bullet']
            pf = p.paragraph_format
            pf.left_indent = Inches(0.25)
            pf.space_after = Pt(2)
        list_items.clear()

    def flush_table():
        nonlocal table_buffer
        if len(table_buffer) >= 2:
            add_table_from_md(doc, table_buffer)
        table_buffer.clear()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip('\n')

        # Code fence
        if stripped.startswith('```'):
            if in_code:
                in_code = False
            else:
                flush_list()
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(stripped)
            try:
                p.style = doc.styles['Code']
            except KeyError:
                p.style = doc.styles['Normal']
            run = p.runs[0] if p.runs else p.add_run(stripped)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            i += 1
            continue

        # Table row
        if stripped.startswith('|') and '|' in stripped[1:]:
            if not in_table:
                flush_list()
                in_table = True
            table_buffer.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                flush_table()
                in_table = False

        # Blank line
        if not stripped.strip():
            flush_list()
            i += 1
            continue

        # Heading
        hm = re.match(r'^(#{1,6})\s+(.+)', stripped)
        if hm:
            flush_list()
            level = len(hm.group(1))
            add_heading(doc, hm.group(2), level)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}$', stripped.strip()):
            p = doc.add_paragraph('─' * 40)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xB5, 0x4A, 0x1C)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            flush_list()
            content = re.sub(r'^>\s*', '', stripped)
            p = add_paragraph_with_inline(doc, content, base_size=11)
            pf = p.paragraph_format
            pf.left_indent = Inches(0.4)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x5A, 0x4A, 0x3A)
            i += 1
            continue

        # List item
        if re.match(r'^(\s*[-*+]|\s*\d+\.)\s+', stripped):
            list_items.append(stripped)
            i += 1
            continue

        # Markdown image: ![alt](url)
        img_m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_m:
            flush_list()
            alt_text = img_m.group(1)
            img_url  = img_m.group(2).strip()
            result = fetch_image(img_url)
            if result:
                _, raw_bytes = result
                try:
                    stream = io.BytesIO(raw_bytes)
                    doc.add_picture(stream, width=Inches(4.25))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if alt_text:
                        cap = doc.add_paragraph(alt_text)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in cap.runs:
                            run.italic = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0x7A, 0x6A, 0x5A)
                except Exception as e:
                    print(f"    [warn] Could not embed image in DOCX: {e}")
                    doc.add_paragraph(f'[Image: {alt_text}]')
            else:
                doc.add_paragraph(f'[Image: {alt_text}]')
            i += 1
            continue

        # Regular paragraph
        flush_list()
        add_paragraph_with_inline(doc, stripped, base_size=12)
        i += 1

    flush_list()
    if in_table:
        flush_table()


# ══════════════════════════════════════════════════════════════════════
#  DOCX GENERATION
# ══════════════════════════════════════════════════════════════════════

def generate_docx():
    print("Generating DOCX...")
    doc = Document()

    # Page setup: 8.5"×11" letter
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = Pt(17)

    # Title page
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run('The Opportunity Mindset')
    run.font.name = 'Georgia'
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1C, 0x25, 0x33)

    doc.add_paragraph()

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run('How to Recognize, Create, and Capture Opportunity')
    sr.font.name = 'Georgia'
    sr.font.size = Pt(14)
    sr.italic = True
    sr.font.color.rgb = RGBColor(0x7A, 0x66, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = ap.add_run('Bernard Baah')
    ar.font.name = 'Georgia'
    ar.font.size = Pt(16)
    ar.font.color.rgb = RGBColor(0xB5, 0x4A, 0x1C)

    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = pp.add_run('Filly Coder · AI Future Series')
    pr.font.name = 'Georgia'
    pr.font.size = Pt(10)
    pr.font.color.rgb = RGBColor(0x9A, 0x80, 0x70)

    doc.add_page_break()

    # Process all files
    for idx, filepath in enumerate(BOOK_FILES):
        if not os.path.exists(filepath):
            print(f"  Skipping missing: {filepath}")
            continue
        print(f"  Adding: {filepath}")
        process_file_to_docx(doc, filepath, is_first=(idx == 0))

    out_path = OUTPUT_DIR / "the-opportunity-mindset.docx"
    doc.save(str(out_path))
    size = out_path.stat().st_size / (1024 * 1024)
    print(f"DOCX saved: {out_path} ({size:.1f} MB)")


# ══════════════════════════════════════════════════════════════════════
#  PRINT HTML GENERATION
# ══════════════════════════════════════════════════════════════════════

def md_to_html_block(text):
    """Very lightweight markdown-to-HTML for the print export."""
    import html as html_lib

    lines = text.split('\n')
    out = []
    in_table = False
    in_code = False
    in_list = False
    in_blockquote = False
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if len(table_rows) < 2:
            table_rows.clear(); return
        out.append('<table>')
        out.append('<thead><tr>')
        headers = [c.strip() for c in table_rows[0].strip().strip('|').split('|')]
        for h in headers:
            out.append(f'<th>{inline_md(h)}</th>')
        out.append('</tr></thead><tbody>')
        for row in table_rows[2:]:
            if not row.strip() or re.match(r'^\|[-| :]+\|$', row.strip()):
                continue
            cells = [c.strip() for c in row.strip().strip('|').split('|')]
            out.append('<tr>')
            for c in cells:
                out.append(f'<td>{inline_md(c)}</td>')
            out.append('</tr>')
        out.append('</tbody></table>')
        table_rows.clear()

    def inline_md(s):
        s = html_lib.escape(s)
        s = re.sub(r'\*\*\*(.+?)\*\*\*', r'<strong><em>\1</em></strong>', s)
        s = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', s)
        s = re.sub(r'\*(.+?)\*', r'<em>\1</em>', s)
        s = re.sub(r'`(.+?)`', r'<code>\1</code>', s)
        s = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', s)
        return s

    for line in lines:
        stripped = line.rstrip()

        if stripped.startswith('```'):
            if in_code:
                out.append('</code></pre>'); in_code = False
            else:
                if in_list: out.append('</ul>'); in_list = False
                flush_table(); in_table = False
                lang = stripped[3:].strip() or ''
                out.append(f'<pre><code class="language-{lang}">')
                in_code = True
            continue

        if in_code:
            out.append(html.escape(stripped) if stripped else '')
            continue

        if stripped.startswith('|') and '|' in stripped[1:]:
            if in_list: out.append('</ul>'); in_list = False
            if in_blockquote: out.append('</blockquote>'); in_blockquote = False
            in_table = True
            table_rows.append(stripped)
            continue
        else:
            if in_table:
                flush_table(); in_table = False

        if not stripped:
            if in_list: out.append('</ul>'); in_list = False
            if in_blockquote: out.append('</blockquote>'); in_blockquote = False
            continue

        hm = re.match(r'^(#{1,6})\s+(.+)', stripped)
        if hm:
            if in_list: out.append('</ul>'); in_list = False
            level = len(hm.group(1))
            out.append(f'<h{level}>{inline_md(hm.group(2))}</h{level}>')
            continue

        if re.match(r'^[-*_]{3,}$', stripped.strip()):
            out.append('<hr/>')
            continue

        if stripped.startswith('>'):
            content = re.sub(r'^>\s*', '', stripped)
            if not in_blockquote:
                out.append('<blockquote>')
                in_blockquote = True
            out.append(f'<p>{inline_md(content)}</p>')
            continue
        else:
            if in_blockquote: out.append('</blockquote>'); in_blockquote = False

        if re.match(r'^(\s*[-*+]|\s*\d+\.)\s+', stripped):
            content = re.sub(r'^(\s*[-*+]|\s*\d+\.)\s+', '', stripped)
            if not in_list:
                out.append('<ul>'); in_list = True
            out.append(f'<li>{inline_md(content)}</li>')
            continue
        else:
            if in_list: out.append('</ul>'); in_list = False

        # Markdown image: ![alt](url)
        img_m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_m:
            alt = html_lib.escape(img_m.group(1))
            url = img_m.group(2).strip()
            result = fetch_image(url)
            if result:
                data_uri, _ = result
                out.append(
                    f'<div class="book-image" style="background-image:url(\'{data_uri}\')" '
                    f'role="img" aria-label="{alt}"></div>'
                )
                if alt:
                    out.append(f'<p class="img-caption">{alt}</p>')
            else:
                out.append(f'<p class="img-placeholder">[Image: {alt}]</p>')
            continue

        out.append(f'<p>{inline_md(stripped)}</p>')

    if in_list: out.append('</ul>')
    if in_blockquote: out.append('</blockquote>')
    if in_code: out.append('</code></pre>')
    if in_table: flush_table()

    return '\n'.join(out)


def generate_print_html():
    print("Generating print HTML...")

    sections_html = []
    for filepath in BOOK_FILES:
        if not os.path.exists(filepath):
            print(f"  Skipping missing: {filepath}")
            continue
        print(f"  Processing: {filepath}")
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        sections_html.append(f'<section class="book-section">\n{md_to_html_block(content)}\n</section>')

    body_content = '\n'.join(sections_html)

    html_doc = f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<title>The Opportunity Mindset — Bernard Baah</title>
<link href="https://fonts.googleapis.com/css2?family=Lora:ital,wght@0,400;0,500;0,600;0,700;1,400;1,500&family=Inter:wght@300;400;500;600&display=swap" rel="stylesheet"/>
<style>
/* ── Screen styles ─────────────────────────────────────── */
* {{ box-sizing: border-box; margin: 0; padding: 0; }}

body {{
  font-family: 'Lora', Georgia, serif;
  font-size: 11.5pt;
  line-height: 1.75;
  color: #1a1a18;
  background: #f5f2ed;
  padding: 40px 20px;
}}

.print-toolbar {{
  position: fixed;
  top: 20px; right: 20px;
  z-index: 1000;
  display: flex; gap: 10px;
}}

.print-toolbar button {{
  padding: 10px 20px;
  background: #B54A1C; color: #fff;
  border: none; border-radius: 4px;
  font-family: 'Inter', sans-serif;
  font-size: 13px; font-weight: 600;
  cursor: pointer; box-shadow: 0 4px 16px rgba(0,0,0,.2);
  transition: opacity .2s;
}}
.print-toolbar button:hover {{ opacity: .85; }}
.print-toolbar .note {{
  background: rgba(0,0,0,.65);
  color: #fff; padding: 10px 16px;
  border-radius: 4px; font-size: 11px;
  font-family: 'Inter', sans-serif;
  line-height: 1.5; max-width: 220px;
  display: none;
}}
.print-toolbar:hover .note {{ display: block; }}

.book-page {{
  max-width: 640px;
  margin: 0 auto;
  background: #fff;
  box-shadow: 0 2px 32px rgba(0,0,0,.08);
  padding: 80px 80px 100px;
}}

/* Title page */
.title-page {{
  min-height: 600px;
  display: flex; flex-direction: column;
  align-items: center; justify-content: center;
  text-align: center;
  border-bottom: 2px solid #B54A1C;
  margin-bottom: 60px;
  page-break-after: always;
}}
.title-page .main-title {{
  font-family: 'Lora', serif;
  font-size: 32pt; font-weight: 700;
  color: #1C2533; line-height: 1.15;
  margin-bottom: 20px;
}}
.title-page .subtitle {{
  font-size: 13pt; font-style: italic;
  color: #7A6655; margin-bottom: 40px;
}}
.title-page .author {{
  font-size: 15pt; font-weight: 600;
  color: #B54A1C; letter-spacing: .04em;
}}
.title-page .publisher {{
  font-family: 'Inter', sans-serif;
  font-size: 9pt; letter-spacing: .2em;
  text-transform: uppercase; color: #9A8070;
  margin-top: 8px;
}}

/* Sections */
.book-section {{
  page-break-before: always;
  padding-top: 40px;
}}
.book-section:first-of-type {{ page-break-before: avoid; }}

/* Typography */
h1 {{ font-size: 22pt; font-weight: 700; color: #1C2533; margin: 48px 0 16px; line-height: 1.2; page-break-after: avoid; }}
h2 {{ font-size: 15pt; font-weight: 600; color: #1C2533; margin: 36px 0 12px; line-height: 1.3; page-break-after: avoid; }}
h3 {{ font-size: 12.5pt; font-weight: 600; color: #3A4A5A; margin: 28px 0 8px; page-break-after: avoid; }}
h4 {{ font-size: 11.5pt; font-weight: 600; color: #5A6A7A; margin: 20px 0 6px; page-break-after: avoid; }}
h5, h6 {{ font-size: 11pt; font-style: italic; margin: 16px 0 4px; }}

p {{ margin-bottom: 10pt; text-align: justify; }}
p + p {{ text-indent: 1.5em; }}
h1 + p, h2 + p, h3 + p, blockquote + p {{ text-indent: 0; }}

strong {{ font-weight: 600; }}
em {{ font-style: italic; }}
code {{ font-family: 'Courier New', monospace; font-size: 9.5pt; background: #f0ece6; padding: 1px 4px; border-radius: 2px; }}
pre {{ background: #f0ece6; padding: 14px 16px; margin: 16px 0; border-left: 3px solid #B54A1C; overflow-x: auto; }}
pre code {{ background: none; padding: 0; font-size: 9pt; line-height: 1.6; }}

blockquote {{
  border-left: 3px solid #B54A1C;
  margin: 20px 0; padding: 12px 20px;
  color: #5A4A3A; font-style: italic;
  background: #faf7f3;
}}
blockquote p {{ text-indent: 0; margin-bottom: 6pt; }}

ul, ol {{ margin: 10pt 0 10pt 24pt; }}
li {{ margin-bottom: 4pt; }}

hr {{ border: none; border-top: 1px solid #D0C0B0; margin: 32px 0; }}

/* Images */
.book-image {{
  width: 100%;
  height: 280px;
  background-size: cover;
  background-position: center;
  background-repeat: no-repeat;
  border-radius: 4px;
  margin: 20px 0 6px;
  page-break-inside: avoid;
}}
.img-caption {{
  text-align: center;
  font-style: italic;
  font-size: 9.5pt;
  color: #7A6655;
  margin-bottom: 16pt;
  text-indent: 0 !important;
}}
.img-placeholder {{
  text-align: center;
  color: #aaa;
  font-style: italic;
  font-size: 9pt;
  padding: 20px;
  border: 1px dashed #ddd;
  margin: 16px 0;
}}

table {{ width: 100%; border-collapse: collapse; margin: 20px 0; font-size: 10pt; page-break-inside: avoid; }}
th {{ background: #1C2533; color: #F0E8D8; padding: 8px 10px; text-align: left; font-family: 'Inter', sans-serif; font-weight: 600; font-size: 9.5pt; }}
td {{ padding: 7px 10px; border-bottom: 1px solid #E0D8D0; vertical-align: top; }}
tr:nth-child(even) td {{ background: #faf7f3; }}

/* ── Print styles ─────────────────────────────────────── */
@media print {{
  * {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}

  body {{ background: #fff; padding: 0; }}
  .print-toolbar {{ display: none !important; }}
  .book-page {{
    max-width: none; box-shadow: none;
    padding: 0;
  }}

  @page {{
    size: 8.5in 11in;
    margin: 1in;
    @top-center {{ content: "The Opportunity Mindset"; font-size: 9pt; color: #9A8070; }}
    @bottom-center {{ content: counter(page); font-size: 9pt; }}
  }}

  .title-page {{
    min-height: 90vh;
    page-break-after: always;
  }}

  .book-section {{
    page-break-before: always;
  }}

  h1, h2, h3 {{ page-break-after: avoid; }}
  table, figure {{ page-break-inside: avoid; }}
  blockquote {{ page-break-inside: avoid; }}
}}
</style>
</head>
<body>

<div class="print-toolbar">
  <div class="note">
    Click "Print / Save PDF" then choose<br>
    <strong>Save as PDF</strong> in your printer dialog.<br>
    Set paper size to <strong>6×9 inches</strong>.
  </div>
  <button onclick="window.print()">🖨 Print / Save PDF</button>
</div>

<div class="book-page">

  <!-- Title page -->
  <div class="title-page">
    <div class="main-title">The<br>Opportunity<br>Mindset</div>
    <div class="subtitle">How to Recognize, Create, and Capture Opportunity</div>
    <div class="author">Bernard Baah</div>
    <div class="publisher">Filly Coder &middot; AI Future Series</div>
  </div>

  {body_content}

</div>

</body>
</html>'''

    out_path = OUTPUT_DIR / "the-opportunity-mindset-print.html"
    with open(str(out_path), 'w', encoding='utf-8') as f:
        f.write(html_doc)
    size = out_path.stat().st_size / (1024 * 1024)
    print(f"Print HTML saved: {out_path} ({size:.1f} MB)")


# ══════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    import html
    generate_docx()
    generate_print_html()
    print("\nDone! Files are in book/downloads/")
